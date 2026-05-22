import os
import json
import pandas as pd
from docx import Document
from docx.shared import Inches
import argparse
from pathlib import Path
import tempfile
from datetime import datetime, date

# --- Новые импорты для PDF ---
import pdfplumber
import fitz  # PyMuPDF
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4, landscape
from reportlab.platypus import Table, TableStyle, SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.fonts import addMapping

# Проверка и установки зависимостей
def check_dependencies():
    try:
        import pandas as pd
        import openpyxl
        import xlrd
        from docx import Document
        import pdfplumber
        import fitz
        import reportlab
    except ImportError as e:
        print(f"Ошибка: отсутствует библиотека: {e.name}")
        print("Установите необходимые зависимости командой:")
        print("pip install pandas openpyxl xlrd python-docx pdfplumber PyMuPDF reportlab")
        exit(1)

class DateEncoder(json.JSONEncoder):
    def default(self, obj):
        if isinstance(obj, (datetime, date)):
            return obj.isoformat()
        return super().default(obj)

class DocumentContent:
    """
    Класс для хранения извлеченного содержимого: текста и таблиц.
    """

    def __init__(self):
        self.text = ""
        self.tables = []  # Список DataFrame'ов

    def add_text(self, text):
        if self.text:
            self.text += "\n" + text
        else:
            self.text = text

    def add_table(self, df):
        self.tables.append(df)

    def has_tables(self):
        return len(self.tables) > 0

    def has_text(self):
        return bool(self.text.strip())

    def to_dict(self):
        """
        Преобразует содержимое в словарь для JSON.
        """
        result = {}
        if self.has_text():
            result['text'] = self.text
        if self.has_tables():
            result['tables'] = [df.to_dict(orient='records') for df in self.tables]
        return result


def extract_content_from_docx(file_path):
    """
    Извлекает текст и таблицы из .docx файла.
    Возвращает объект DocumentContent.
    """
    content = DocumentContent()
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    content.add_text('\n'.join(full_text))

    for table in doc.tables:
        data = []
        for row in table.rows:
            data.append([cell.text for cell in row.cells])
        if data:
            df = pd.DataFrame(data[1:], columns=data[0])  # Первая строка - заголовки
            content.add_table(df)
    return content


def extract_content_from_pdf(file_path):
    """
    Извлекает текст и таблицы из .pdf файла.
    Возвращает объект DocumentContent.
    """
    content = DocumentContent()
    full_text = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            # Извлечение текста
            text = page.extract_text()
            if text:
                full_text.append(text)
            # Извлечение таблиц
            tab = page.extract_tables()
            for table in tab:
                if table:
                    df = pd.DataFrame(table[1:], columns=table[0])  # Первая строка - заголовки
                    content.add_table(df)
    content.add_text('\n'.join(full_text))
    return content

def load_data(source_path, target_format):
    """
    Загружает данные из файла в зависимости от расширения.
    Возвращает объект DocumentContent, pandas.DataFrame, список DataFrame, или путь к файлу (для docx -> pdf).
    """
    ext = Path(source_path).suffix.lower()

    if ext == '.csv':
        df = pd.read_csv(source_path)
        content = DocumentContent()
        content.add_table(df)
        return content
    elif ext in ['.xlsx', '.xls']:
        # Возвращает словарь {sheet_name: DataFrame}
        dfs = pd.read_excel(source_path, sheet_name=None)
        content = DocumentContent()
        for name, df in dfs.items():
            content.add_table(df)
        return content
    elif ext == '.json':
        # Улучшенная обработка JSON
        content = DocumentContent()
        
        try:
            # Пробуем прочитать как табличный JSON (список объектов)
            df = pd.read_json(source_path)
            content.add_table(df)
        except ValueError:
            # Если не получилось, пробуем прочитать как простой JSON
            with open(source_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Если это список - создаем таблицу
            if isinstance(data, list):
                df = pd.DataFrame(data)
                content.add_table(df)
            # Если это словарь - пробуем преобразовать
            elif isinstance(data, dict):
                # Создаем таблицу из ключ-значение
                df = pd.DataFrame(list(data.items()), columns=['Key', 'Value'])
                content.add_table(df)
            else:
                raise ValueError("JSON должен быть списком объектов или словарем")
        
        return content
    elif ext == '.docx':
        if target_format == 'pdf':
            # Для конвертации в PDF возвращаем путь к файлу
            return source_path
        return extract_content_from_docx(source_path)
    elif ext == '.pdf':
        # Разрешаем конвертацию PDF в Word, Excel, JSON и CSV
        if target_format in ['word', 'excel', 'json', 'csv']:
            return extract_content_from_pdf(source_path)
        else:
            raise ValueError("PDF может быть конвертирован только в Word, Excel, JSON или CSV.")
    else:
        raise ValueError(f"Неподдерживаемый формат файла: {ext}")


def docx_to_pdf_reportlab(docx_path, pdf_path):
    """
    Конвертирует .docx в .pdf с использованием reportlab.
    Извлекает текст и таблицы с поддержкой кириллицы.
    """
    # Регистрируем шрифт с кириллицей
    font_name = register_cyrillic_fonts()
    font_bold = font_name + '-Bold' if font_name != 'Helvetica' else 'Helvetica-Bold'
    
    doc = Document(docx_path)
    
    # Создаем PDF документ
    doc_pdf = SimpleDocTemplate(
        str(pdf_path),
        pagesize=A4,
        rightMargin=72,
        leftMargin=72,
        topMargin=72,
        bottomMargin=72
    )
    
    styles = getSampleStyleSheet()
    
    # Создаем стили с кириллическим шрифтом
    styles.add(ParagraphStyle(
        name='CyrillicNormal',
        parent=styles['Normal'],
        fontName=font_name,
        fontSize=11,
        leading=13
    ))
    styles.add(ParagraphStyle(
        name='CyrillicHeading1',
        parent=styles['Heading1'],
        fontName=font_bold,
        fontSize=16,
        leading=18
    ))
    styles.add(ParagraphStyle(
        name='CyrillicHeading2',
        parent=styles['Heading2'],
        fontName=font_bold,
        fontSize=14,
        leading=16
    ))
    
    style_normal = styles['CyrillicNormal']
    style_h1 = styles['CyrillicHeading1']
    style_h2 = styles['CyrillicHeading2']
    
    story = []
    
    # Обрабатываем элементы документа
    for element in doc.element.body:
        if element.tag.endswith('p'):  # Параграф
            text = ''.join([t.text for t in element.xpath('.//w:t') if t.text])
            if text.strip():
                # Определяем стиль (заголовок или обычный текст)
                if element.xpath('.//w:pPr/w:pStyle/@w:val'):
                    style_name = element.xpath('.//w:pPr/w:pStyle/@w:val')[0]
                    if 'Heading1' in style_name or 'Heading 1' in style_name or style_name == '1':
                        story.append(Paragraph(text.strip(), style_h1))
                    elif 'Heading2' in style_name or 'Heading 2' in style_name or style_name == '2':
                        story.append(Paragraph(text.strip(), style_h2))
                    else:
                        story.append(Paragraph(text.strip(), style_normal))
                else:
                    story.append(Paragraph(text.strip(), style_normal))
                story.append(Spacer(1, 6))
                
        elif element.tag.endswith('tbl'):  # Таблица
            # Извлекаем таблицу
            temp_content = extract_content_from_docx(docx_path)
            # Берем последнюю таблицу (текущую)
            if temp_content.tables:
                df = temp_content.tables[-1]
                
                # Подготовка данных таблицы с Paragraph для кириллицы
                table_data = []
                
                # Заголовки
                headers = []
                for col in df.columns:
                    col_str = str(col).strip() if pd.notna(col) else ''
                    headers.append(Paragraph(col_str, style_normal))
                table_data.append(headers)
                
                # Строки
                for _, row in df.iterrows():
                    row_data = []
                    for cell in row:
                        cell_str = str(cell).strip() if pd.notna(cell) else ''
                        row_data.append(Paragraph(cell_str, style_normal))
                    table_data.append(row_data)
                
                # Расчет ширины колонок
                num_cols = len(table_data[0])
                page_width = A4[0] - 144
                col_width = max(100, min((page_width / num_cols) - 10, 250))
                col_widths = [col_width] * num_cols
                
                # Создаем таблицу
                table = Table(table_data, colWidths=col_widths, repeatRows=1)
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.darkblue),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), font_bold),
                    ('FONTSIZE', (0, 0), (-1, 0), 11),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ('FONTSIZE', (0, 1), (-1, -1), 10),
                    ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey]),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ('LEFTPADDING', (0, 0), (-1, -1), 5),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 5),
                    ('TOPPADDING', (0, 0), (-1, -1), 4),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
                ]))
                
                story.append(table)
                story.append(Spacer(1, 12))
    
    # Строим PDF
    doc_pdf.build(story)
    print(f"✓ Конвертировано в PDF: {pdf_path}")

def save_data(data, target_format, output_path):
    """
    Сохраняет данные в нужный формат.
    Возвращает список путей (Path) к созданным файлам.
    """
    dir_path = Path(output_path).parent
    os.makedirs(dir_path, exist_ok=True)

    # Список для отслеживания созданных файлов
    saved_files = []

    if isinstance(data, DocumentContent):
        # Обработка объекта DocumentContent
        if target_format == 'csv':
            if data.has_tables():
                for i, df in enumerate(data.tables):
                    # Имя файла меняется на ..._table_1.csv
                    p = output_path.with_name(f"{output_path.stem}_table_{i+1}.csv")
                    # ИСПРАВЛЕНИЕ: добавляем encoding='utf-8-sig'
                    df.to_csv(p, index=False, encoding='utf-8-sig')
                    saved_files.append(p)  # <-- ЗАПОМИНАЕМ ПУТЬ
                    
                print(f"Сохранено {len(data.tables)} CSV файлов в {dir_path}")
            else:
                print("Файл не содержит таблиц, невозможно сохранить в CSV.")
        
        elif target_format == 'json':
            output_file = output_path.with_suffix('.json')
            json_data = data.to_dict()
            with open(output_file, 'w', encoding='utf-8') as f:
                json.dump(json_data, f, ensure_ascii=False, indent=2)
            print(f"Сохранено в {output_file}")
            saved_files.append(output_file)  # <-- ЗАПОМИНАЕМ ПУТЬ
            
        elif target_format == 'excel':
            if data.has_tables():
                output_file = output_path.with_suffix('.xlsx')
                with pd.ExcelWriter(output_file) as writer:
                    for i, df in enumerate(data.tables):
                        df.to_excel(writer, sheet_name=f'Table_{i+1}', index=False)
                print(f"Сохранено в {output_file}")
                saved_files.append(output_file)  # <-- ЗАПОМИНАЕМ ПУТЬ
            else:
                print("Файл не содержит таблиц, невозможно сохранить в Excel.")
                
        elif target_format == 'word':
            output_file = output_path.with_suffix('.docx')
            doc = Document()
            if data.has_text():
                doc.add_paragraph(data.text)
                if data.has_tables():
                    doc.add_paragraph('')
            if data.has_tables():
                for i, df in enumerate(data.tables):
                    doc.add_heading(f'Таблица {i+1}', level=1)
                    table = doc.add_table(rows=1, cols=len(df.columns))
                    table.style = 'Table Grid'
                    hdr_cells = table.rows[0].cells
                    for j, col_name in enumerate(df.columns):
                        hdr_cells[j].text = str(col_name)
                    for _, row in df.iterrows():
                        row_cells = table.add_row().cells
                        for j, value in enumerate(row):
                            row_cells[j].text = str(value)
            doc.save(output_file)
            print(f"Сохранено в {output_file}")
            saved_files.append(output_file)  # <-- ЗАПОМИНАЕМ ПУТЬ
            
        elif target_format == 'pdf':
            # УЛУЧШЕННАЯ обработка DocumentContent -> PDF
            # save_document_content_to_pdf должна создавать файл output_path.with_suffix('.pdf')
            save_document_content_to_pdf(data, output_path)
            output_file = output_path.with_suffix('.pdf')
            saved_files.append(output_file)  # <-- ЗАПОМИНАЕМ ПУТЬ PDF

    elif isinstance(data, (str, Path)) and target_format == 'pdf':
        # Специальная ветка для пути к файлу Word -> PDF
        docx_path = data
        output_pdf = output_path.with_suffix('.pdf')
        docx_to_pdf_reportlab(docx_path, output_pdf)
        saved_files.append(output_pdf)  # <-- ЗАПОМИНАЕМ ПУТЬ PDF
        
    else:
        print(f"Невозможно конвертировать тип {type(data)} в {target_format}.")

    return saved_files  # <-- ВОЗВРАЩАЕМ СПИСОК ФАЙЛОВ


def register_cyrillic_fonts():
    """Регистрирует шрифты с поддержкой кириллицы."""
    # Пробуем найти доступные шрифты с кириллицей
    font_paths = [
        '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',
        '/usr/share/fonts/TTF/DejaVuSans.ttf',
        'C:/Windows/Fonts/arial.ttf',
        '/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf',
    ]
    
    font_found = False
    for font_path in font_paths:
        if os.path.exists(font_path):
            try:
                pdfmetrics.registerFont(TTFont('DejaVuSans', font_path))
                pdfmetrics.registerFont(TTFont('DejaVuSans-Bold', font_path.replace('Regular', 'Bold').replace('DejaVuSans', 'DejaVuSans-Bold')))
                addMapping('DejaVuSans', 0, 0, 'DejaVuSans')
                addMapping('DejaVuSans', 1, 0, 'DejaVuSans-Bold')
                print(f"✓ Шрифт зарегистрирован: {font_path}")
                font_found = True
                return 'DejaVuSans'
            except Exception as e:
                print(f"Не удалось зарегистрировать шрифт {font_path}: {e}")
                continue
    
    if not font_found:
        print("⚠ Не найден шрифт с кириллицей. Используем стандартный (кириллица может не отображаться)")
        return 'Helvetica'
    
    return 'Helvetica'


def save_document_content_to_pdf(data, output_path):
    """
    Улучшенная функция с правильным применением кириллического шрифта.
    """
    # Регистрируем шрифт
    font_name = register_cyrillic_fonts()
    font_bold = font_name + '-Bold' if font_name != 'Helvetica' else 'Helvetica-Bold'
    
    output_file = output_path.with_suffix('.pdf')
    
    # Определяем ориентацию страницы
    max_cols = 0
    if data.has_tables():
        for df in data.tables:
            max_cols = max(max_cols, len(df.columns))
    
    # Если колонок больше 5, используем ландшафт
    if max_cols > 5:
        page_size = landscape(A4)
        margins = {'left': 30, 'right': 30, 'top': 30, 'bottom': 30}
    else:
        page_size = A4
        margins = {'left': 40, 'right': 40, 'top': 40, 'bottom': 40}
    
    doc = SimpleDocTemplate(
        str(output_file),
        pagesize=page_size,
        leftMargin=margins['left'],
        rightMargin=margins['right'],
        topMargin=margins['top'],
        bottomMargin=margins['bottom']
    )
    
    styles = getSampleStyleSheet()
    
    # Создаем стили для таблицы
    styles.add(ParagraphStyle(
        name='TableHeader',
        parent=styles['Normal'],
        fontName=font_bold,
        fontSize=8,
        textColor=colors.whitesmoke,
        alignment=1,  # center
        spaceBefore=2,
        spaceAfter=2
    ))
    styles.add(ParagraphStyle(
        name='TableCell',
        parent=styles['Normal'],
        fontName=font_name,
        fontSize=7,
        alignment=1,  # center
        spaceBefore=1,
        spaceAfter=1
    ))
    
    styleN = styles['Normal']
    styleH = styles['Heading1']
    styleHeader = styles['TableHeader']
    styleCell = styles['TableCell']
    
    story = []
    
    # Добавляем текст если есть
    if data.has_text():
        for paragraph in data.text.split('\n\n'):
            if paragraph.strip():
                story.append(Paragraph(paragraph.strip(), styleN))
                story.append(Spacer(1, 12))
    
    # Добавляем таблицы
    if data.has_tables():
        for i, df in enumerate(data.tables):
            print(f"Сохраняю таблицу {i+1}: {df.shape}")
            
            # Заголовок таблицы
            story.append(Paragraph(f"Таблица {i+1}", styleH))
            story.append(Spacer(1, 10))
            
            # Подготовка данных таблицы - ОБЯЗАТЕЛЬНО используем Paragraph
            table_data = []
            
            # Заголовки колонок
            headers = []
            for col in df.columns:
                col_str = format_header_value(col)
                headers.append(Paragraph(col_str, styleHeader))
            table_data.append(headers)
            
            # Строки данных
            for _, row in df.iterrows():
                row_data = []
                for cell in row:
                    cell_str = format_cell_value(cell)
                    row_data.append(Paragraph(cell_str, styleCell))
                table_data.append(row_data)
            
            # Расчет ширины колонок
            num_cols = len(table_data[0])
            page_width = page_size[0] - margins['left'] - margins['right']
            
            if num_cols > 8:
                col_width = 85
                font_size = 6
            elif num_cols > 6:
                col_width = 100
                font_size = 7
            elif num_cols > 4:
                col_width = 130
                font_size = 8
            else:
                col_width = max(100, min((page_width / num_cols) - 15, 200))
                font_size = 9
            
            col_widths = [col_width] * num_cols
            
            # Создаем таблицу
            table = Table(table_data, colWidths=col_widths, repeatRows=1)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.darkblue),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), font_bold),
                ('FONTSIZE', (0, 0), (-1, 0), 8),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 6),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
                ('FONTSIZE', (0, 1), (-1, -1), 7),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey]),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('LEFTPADDING', (0, 0), (-1, -1), 3),
                ('RIGHTPADDING', (0, 0), (-1, -1), 3),
                ('TOPPADDING', (0, 0), (-1, -1), 2),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
            ]))
            
            story.append(table)
            story.append(Spacer(1, 20))
    
    # Строим PDF
    try:
        doc.build(story)
        print(f"✓ Сохранено в {output_file}")
        print(f"  Ориентация: {'landscape' if max_cols > 5 else 'portrait'}")
        print(f"  Максимум колонок: {max_cols}")
        print(f"  Шрифт: {font_name}")
    except Exception as e:
        print(f"✗ Ошибка при создании PDF: {e}")
        import traceback
        traceback.print_exc()
        raise


def format_header_value(value):
    """Форматирует значение заголовка."""
    if isinstance(value, (pd.Timestamp, datetime, date)):
        return value.strftime('%d.%m.%Y')
    
    col_str = str(value).strip()
    
    if 'Unnamed' in col_str or col_str == '':
        return ''
    
    # Форматирование дат в строках, если вдруг они попали сюда
    if len(col_str) > 10 and ('-' in col_str or ':' in col_str):
        try:
            if ' ' in col_str:
                date_part = col_str.split(' ')[0]
            else:
                date_part = col_str
            
            if len(date_part) >= 10:
                parts = date_part.split('-')
                if len(parts) == 3:
                    return f"{parts[2]}.{parts[1]}.{parts[0][:4]}"
        except:
            pass
    
    # УБРАНО: Ограничение на длину заголовка
    # if len(col_str) > 15:
    #     return col_str[:13] + '..'
    
    return col_str

def format_cell_value(value):
    """Форматирует значение ячейки."""
    
    # 1. ИСПРАВЛЕНИЕ: Если это список или словарь (часто бывает в JSON),
    # сразу конвертируем в строку, чтобы избежать ошибок pandas.
    if isinstance(value, (list, dict)):
        return str(value)

    # 2. Проверка на NaN (теперь безопасна для простых типов)
    try:
        if pd.isna(value):
            return ''
    except (ValueError, TypeError):
        # На случай если значение сложное
        pass

    # 3. Если дата/время
    if isinstance(value, (pd.Timestamp, datetime, date)):
        return value.strftime('%d.%m.%Y %H:%M')

    # 4. Преобразуем в строку
    cell_str = str(value).strip()

    # 5. Если это дата в строковом формате (например "2024-01-01")
    if len(cell_str) > 10 and ('-' in cell_str or ':' in cell_str):
        try:
            if ' ' in cell_str:
                date_time_parts = cell_str.split(' ')
                date_part = date_time_parts[0]
                time_part = date_time_parts[1] if len(date_time_parts) > 1 else '00:00'
                
                if len(date_part) >= 10:
                    parts = date_part.split('-')
                    if len(parts) == 3:
                        return f"{parts[2]}.{parts[1]}.{parts[0][:4]}"
        except:
            pass
    
    # 6. Ограничение длины
    if len(cell_str) > 30:
        return cell_str[:28] + '..'
    
    return cell_str

def main():
    parser = argparse.ArgumentParser(description='Конвертер форматов данных (Excel, CSV, JSON, Word, PDF)')
    parser.add_argument('input_file', type=str, help='Путь к исходному файлу')
    parser.add_argument('output_format', type=str, choices=['csv', 'json', 'excel', 'word', 'pdf'],
                        help='Целевой формат')
    parser.add_argument('output_folder', type=str, help='Папка для сохранения результата')

    args = parser.parse_args()

    check_dependencies()

    input_path = Path(args.input_file)
    output_folder = Path(args.output_folder)
    target_format = args.output_format

    if not input_path.exists():
        print(f"Ошибка: файл не найден - {input_path}")
        return

    try:
        print(f"Загрузка данных из {input_path}...")
        data = load_data(input_path, target_format)

        output_name = input_path.stem
        output_path = output_folder / output_name

        print(f"Конвертация в формат {target_format}...")
        save_data(data, target_format, output_path)

    except ValueError as e:
        print(f"Ошибка: {e}")
    except Exception as e:
        print(f"Произошла ошибка: {e}")


if __name__ == '__main__':
    main()
